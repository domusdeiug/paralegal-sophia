import re
import io
from datetime import datetime, timezone
from google.adk.tools import ToolContext
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from ..lib.supabase_client import get_supabase


# ---------------------------------------------------------------------------
# Inline parser: **bold**, *italic*, plain text
# ---------------------------------------------------------------------------

def _parse_inline(line: str) -> list[dict]:
    """Return list of {text, bold, italic} dicts."""
    runs = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))')
    for m in pattern.finditer(line):
        if m.group(2):
            runs.append({"text": m.group(2), "bold": True, "italic": False})
        elif m.group(3):
            runs.append({"text": m.group(3), "bold": False, "italic": True})
        elif m.group(4):
            runs.append({"text": m.group(4), "bold": False, "italic": False})
    return runs or [{"text": line, "bold": False, "italic": False}]


def _add_runs(para, runs: list[dict]):
    for r in runs:
        run = para.add_run(r["text"])
        run.bold = r.get("bold", False)
        run.italic = r.get("italic", False)


# ---------------------------------------------------------------------------
# Markdown → python-docx Document
# ---------------------------------------------------------------------------

def _markdown_to_docx(markdown: str) -> Document:
    doc = Document()

    # Page layout: A4, 1-inch margins
    section = doc.sections[0]
    section.page_width = int(11906 * 914.4 / 12240)   # twips→EMU approximation
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Default style: Times New Roman 12pt
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    lines = markdown.split("\n")
    i = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Heading 1
        m = re.match(r"^# (.+)", line)
        if m:
            p = doc.add_heading(level=1)
            p.clear()
            run = p.add_run(m.group(1))
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = "Times New Roman"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Heading 2
        m = re.match(r"^## (.+)", line)
        if m:
            p = doc.add_heading(level=2)
            p.clear()
            run = p.add_run(m.group(1))
            run.bold = True
            run.underline = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
            i += 1
            continue

        # Heading 3
        m = re.match(r"^### (.+)", line)
        if m:
            p = doc.add_heading(level=3)
            p.clear()
            run = p.add_run(m.group(1))
            run.bold = True
            run.italic = True
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            content = re.sub(r"^>\s?", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            _add_runs(p, _parse_inline(content))
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)(\d+)\. (.+)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, _parse_inline(m.group(3)))
            i += 1
            continue

        # Bullet list
        m = re.match(r"^(\s*)[-*] (.+)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, _parse_inline(m.group(2)))
            i += 1
            continue

        # Table block
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            # Filter separator rows
            data_rows = [l for l in table_lines if not re.match(r"^\|[\s|:-]+\|$", l)]
            if data_rows:
                parsed = []
                for row_line in data_rows:
                    cells = [c.strip() for c in row_line.strip("|").split("|")]
                    parsed.append(cells)
                col_count = max(len(r) for r in parsed)
                tbl = doc.add_table(rows=len(parsed), cols=col_count)
                tbl.style = "Table Grid"
                for ri, row_data in enumerate(parsed):
                    row = tbl.rows[ri]
                    for ci, cell_text in enumerate(row_data):
                        if ci < col_count:
                            cell = row.cells[ci]
                            cell.text = cell_text
                            if ri == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
            doc.add_paragraph()
            continue

        # Empty line
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_runs(p, _parse_inline(line))
        i += 1

    return doc


# ---------------------------------------------------------------------------
# Tool
# ---------------------------------------------------------------------------

def generate_docx(markdown_content: str, filename: str, tool_context: ToolContext) -> str:
    """Draft a formal Ugandan court submission as a Word document and return a download URL.
    Only call this when the user explicitly wants a document drafted and saved.
    """
    user_id: str = tool_context.state["user_id"]
    supabase = get_supabase()

    # Convert markdown → docx bytes
    doc = _markdown_to_docx(markdown_content)
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    # Ensure filename ends with .docx
    if not filename.endswith(".docx"):
        filename = filename + ".docx"

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S")
    safe_name = re.sub(r"[^\w.-]", "_", filename)
    storage_path = f"{user_id}/court_submissions/{timestamp}_{safe_name}"

    # Upload to Supabase Storage
    supabase.storage.from_("user-documents").upload(
        storage_path,
        docx_bytes,
        file_options={
            "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "upsert": "false",
        },
    )

    # Insert record into user_documents
    supabase.table("user_documents").insert({
        "user_id": user_id,
        "kind": "court_submission",
        "file_name": filename,
        "storage_path": storage_path,
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "size_bytes": len(docx_bytes),
    }).execute()

    # Create 7-day signed URL
    signed = supabase.storage.from_("user-documents").create_signed_url(
        storage_path, 60 * 60 * 24 * 7
    )

    return signed["signedURL"]
